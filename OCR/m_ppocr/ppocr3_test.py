#!/usr/bin/env python
# -*- coding: utf-8 -*-
'''
@File    :   ppocr3_test.py
@Time    :   2024/05/20 10:00:39
@Author  :   hgh 
@Version :   1.0
@Desc    :    
'''

# import module


import os
import cv2 as cv
import numpy as np
import fastdeploy as fd

from docx import Document
from docx.shared import Inches

# 检测是否包含数字
def contains_digit(s):
    return any(c.isdigit() for c in s)

# 检测是否包含数字，且判断数字长度
def check_digit_length(s):
    digit_count = 0
    continuous_digit_start = -1

    for index, c in enumerate(s):
        if c.isdigit():
            if continuous_digit_start == -1:
                continuous_digit_start = index
            digit_count += 1
        else:
            if continuous_digit_start != -1:
                if digit_count > 0:  # 如果找到了至少一个数字
                    num_str = s[continuous_digit_start:continuous_digit_start + digit_count]
                    num_len = len(num_str)
                    print(f"找到数字：{num_str}，长度为：{num_len}")
                continuous_digit_start = -1
                digit_count = 0

    # 检查末尾是否有连续的数字
    if continuous_digit_start != -1 and digit_count > 0:
        num_str = s[continuous_digit_start:]
        num_len = len(num_str)
        print(f"找到数字：{num_str}，长度为：{num_len}")
        
        return  num_str, num_len

# 写入文档到word
def write_to_word(txt):
    # 创建一个新的Word文档对象
    doc = Document()

    doc.add_paragraph(txt)
        
# 检测是否包含数字，返回长度
def digit_length(s):
    digit_count = 0
    current_count = 0
    for c in s:
        if c.isdigit():
            current_count += 1
        elif current_count > 0:
            digit_count += current_count
            current_count = 0
    if current_count > 0:
        digit_count += current_count
    return digit_count


######################################## ppocr 部分 ##################################################

cls_bs = 1
rec_bs = 6
device = 'cpu'

rec_model = r'../data/weights/ch_PP-OCRv3_rec_infer'
rec_label_file = r'../data/weights/ppocr_keys_v1.txt'

det_model = r'../data/weights/ch_PP-OCRv3_det_infer'

rec_model_file = os.path.join(rec_model, "inference.pdmodel")
rec_params_file = os.path.join(rec_model, "inference.pdiparams")

det_model_file = os.path.join(det_model, "inference.pdmodel")
det_params_file = os.path.join(det_model, "inference.pdiparams")



det_option = fd.RuntimeOption()
# cls_option = fd.RuntimeOption()
rec_option = fd.RuntimeOption()

det_option.use_ort_backend()
# cls_option.use_ort_backend()
rec_option.use_ort_backend()


rec_model = fd.vision.ocr.Recognizer(
    rec_model_file, rec_params_file, rec_label_file, runtime_option=rec_option)

det_model = fd.vision.ocr.DBDetector(
    det_model_file, det_params_file, runtime_option=det_option)


# Parameters settings for pre and post processing of Det/Cls/Rec Models.
# All parameters are set to default values.
det_model.preprocessor.max_side_len = 960
det_model.postprocessor.det_db_thresh = 0.3
det_model.postprocessor.det_db_box_thresh = 0.6
det_model.postprocessor.det_db_unclip_ratio = 1.5
det_model.postprocessor.det_db_score_mode = "slow"
det_model.postprocessor.use_dilation = False
# cls_model.postprocessor.cls_thresh = 0.9

# det_model.preprocessor.max_side_len = 960
# det_model.postprocessor.det_db_thresh = 0.3
# det_model.postprocessor.det_db_box_thresh = 0.6
# det_model.postprocessor.det_db_unclip_ratio = 1.5
# det_model.postprocessor.det_db_score_mode = "slow"  # "middle"
# det_model.postprocessor.use_dilation = False
# # cls_model.postprocessor.cls_thresh = 0.9

# Create PP-OCRv3, if cls_model is not needed, just set cls_model=None .
ppocr_v3 = fd.vision.ocr.PPOCRv3(
    det_model=det_model, rec_model=rec_model)

ppocr_v3.cls_batch_size = cls_bs
ppocr_v3.rec_batch_size = rec_bs


# img_path ='data/images/12.jpg'
# img_path ='./src/test.jpg'
img_path = r'../data/images/number01.png'
# Read the input image
im = cv.imread(img_path)

# Predict and reutrn the results
result = ppocr_v3.predict(im)
# print(result)


# ##########################  内容写入word文档  ######################### 
# doc = Document()
# boxes = []
# for index, text in enumerate(result.text):
#     print(text)
#     doc.add_paragraph(text)

# # 保存文档
# doc.save('example.docx')



# # 接口绘制内容  Visuliaze the results.
# vis_im = fd.vision.vis_ppocr(im, result)
# cv.imwrite("../data/results/result.jpg", vis_im)
# print("Visualized result save in ./result.jpg")



text_len = 3
if len(result.text) > 0:
    for index, text in enumerate(result.text):
        if contains_digit(text):
            num_len = digit_length(text)
            if (num_len) < text_len:
                    print(index, text)
            
                    # # text = keep_digits_and_letters(text)
                    # contect += text
                    # ui.textEdit_result.append(text)
                    box = result.boxes[index]       # 4个坐标点，顺序为左下，右下，右上，左上
                    # left_up = box[0], box[1]-5
                    # right_down = box[4], box[5]+10
                    # print('*'*30)
                    # print(box)
                    # print('-'*30)
                    
                    left_up = box[6], box[7]
                    right_down = box[2], box[3]
                    # 生成一个随机的BGR颜色
                    random_color = (np.random.randint(0, 256), np.random.randint(0, 256), np.random.randint(0, 256))

                    cv.rectangle(im, left_up, right_down, random_color, thickness=2, lineType=cv.LINE_AA)



cv.imwrite("../data/results/result.jpg", im)
print("Visualized result save in ./result.jpg")







##########################  绘制检测框  ######################### 
# boxes = []
# for index, text in enumerate(result.text):
#     print(text)
#     doc.add_paragraph(text)
    # print(type(text))
#     if contains_digit(text):
#             num_len = digit_length(text)
#         # if (num_len) > 2:
#             print(index, text)
#             box = result.boxes[index]
#             left_up = box[0], box[1]
#             right_up = box[2], box[3]
#             right_down = box[4], box[5]
#             left_down = box[6], box[7]
#             print(left_up)
#             print(right_down)
#             print('*'*30)
#             boxes.append([left_up, right_down])

# # print(boxes)
# # print(len(boxes))

# for box in boxes:
#     # print(box)
    
#     cv.rectangle(im, box[0], box[1], (0, 0, 255))

# cv.imwrite('res.png', im)
# print('successful!')
